from docx import Document
from app.utils.logging_config import logger

def read_docx(file_path: str) -> str:
    """Extract text content from a DOCX file, including paragraphs and tables."""
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        # Include table cells content
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error reading DOCX {file_path}: {e}", exc_info=True)
        return f"Error reading DOCX: {e}"
