import os
import re
from typing import List, Dict, Any, Optional
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

from app.utils.logging_config import logger

# Import refactored tools from app/tools
from app.tools.pdf_reader import read_pdf
from app.tools.csv_reader import analyze_csv
from app.tools.docx_reader import read_docx
from app.tools.web_search import get_web_search_provider, BaseWebSearch, TavilySearch, MockSearch
from app.tools.calculator import calculate

def convert_md_to_html_formatting(text: str) -> str:
    """Helper to convert markdown formatting to ReportLab-friendly HTML tags."""
    # Convert bold **text** to <b>text</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Convert italic *text* to <i>text</i>
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    # Escape ampersands but preserve HTML tags
    text = text.replace("&", "&amp;")
    # Fix back any accidental double escapes of tags if they exist
    text = text.replace("&amp;lt;", "&lt;").replace("&amp;gt;", "&gt;")
    return text

def generate_pdf_report(report_markdown: str, output_path: str, title: str = "Research Report") -> None:
    """Compile Markdown report into a professional PDF using ReportLab."""
    try:
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles for a premium design
        title_style = ParagraphStyle(
            name='ReportTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=26,
            leading=30,
            textColor=colors.HexColor('#1E293B'), # Slate 800
            spaceAfter=15
        )
        
        subtitle_style = ParagraphStyle(
            name='ReportSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=12,
            leading=16,
            textColor=colors.HexColor('#64748B'), # Slate 500
            spaceAfter=25
        )
        
        h1_style = ParagraphStyle(
            name='ReportH1',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            textColor=colors.HexColor('#0F172A'), # Slate 900
            spaceBefore=15,
            spaceAfter=8,
            keepWithNext=True
        )

        h2_style = ParagraphStyle(
            name='ReportH2',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=13,
            leading=17,
            textColor=colors.HexColor('#0F172A'),
            spaceBefore=10,
            spaceAfter=6,
            keepWithNext=True
        )
        
        body_style = ParagraphStyle(
            name='ReportBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#334155'), # Slate 700
            spaceAfter=10
        )
        
        bullet_style = ParagraphStyle(
            name='ReportBullet',
            parent=body_style,
            leftIndent=20,
            firstLineIndent=-10,
            spaceAfter=5
        )

        story = []
        
        # Cover / Header Title
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')} | Multi-Agent AI Research System", subtitle_style))
        story.append(Spacer(1, 0.15 * inch))
        
        # Parse Markdown line by line
        lines = report_markdown.split('\n')
        
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
                
            # Headers
            if line_str.startswith('# '):
                story.append(Spacer(1, 0.1 * inch))
                header_text = convert_md_to_html_formatting(line_str[2:])
                story.append(Paragraph(header_text, h1_style))
            elif line_str.startswith('## '):
                story.append(Spacer(1, 0.05 * inch))
                header_text = convert_md_to_html_formatting(line_str[3:])
                story.append(Paragraph(header_text, h2_style))
            elif line_str.startswith('### '):
                header_text = convert_md_to_html_formatting(line_str[4:])
                story.append(Paragraph(header_text, h2_style))
            # Bullets
            elif line_str.startswith('* ') or line_str.startswith('- '):
                bullet_text = convert_md_to_html_formatting(line_str[2:])
                story.append(Paragraph(f"&bull; {bullet_text}", bullet_style))
            # Numbered lists
            elif re.match(r'^\d+\.\s', line_str):
                match = re.match(r'^(\d+)\.\s(.*)', line_str)
                num = match.group(1)
                item_text = convert_md_to_html_formatting(match.group(2))
                story.append(Paragraph(f"{num}. {item_text}", bullet_style))
            # Standard Paragraph
            else:
                body_text = convert_md_to_html_formatting(line_str)
                story.append(Paragraph(body_text, body_style))
                
        # Build PDF
        doc.build(story)
        logger.info(f"PDF successfully generated at {output_path}")
    except Exception as e:
        logger.error(f"Failed to generate PDF: {e}", exc_info=True)
        raise e

def generate_docx_report(report_markdown: str, output_path: str, title: str = "Research Report") -> None:
    """Compile Markdown report into a professional DOCX file using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    
    try:
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(title)
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        # Subtitle
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(f"Generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')} | Multi-Agent AI Research System")
        sub_run.font.name = 'Arial'
        sub_run.font.size = Pt(10)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        
        lines = report_markdown.split('\n')
        
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
                
            # Headings
            if line_str.startswith('# '):
                p = doc.add_paragraph()
                run = p.add_run(line_str[2:])
                run.font.name = 'Arial'
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            elif line_str.startswith('## '):
                p = doc.add_paragraph()
                run = p.add_run(line_str[3:])
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
            elif line_str.startswith('### '):
                p = doc.add_paragraph()
                run = p.add_run(line_str[4:])
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
            # Bullets
            elif line_str.startswith('* ') or line_str.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                parts = re.split(r'\*\*(.*?)\*\*', line_str[2:])
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    if idx % 2 == 1:
                        run.bold = True
            # Numbered Lists
            elif re.match(r'^\d+\.\s', line_str):
                match = re.match(r'^(\d+)\.\s(.*)', line_str)
                num = match.group(1)
                item_text = match.group(2)
                p = doc.add_paragraph(style='List Number')
                parts = re.split(r'\*\*(.*?)\*\*', item_text)
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    if idx % 2 == 1:
                        run.bold = True
            # Standard Paragraph
            else:
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.*?)\*\*', line_str)
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                    if idx % 2 == 1:
                        run.bold = True
                p.paragraph_format.space_after = Pt(6)
                
        doc.save(output_path)
        logger.info(f"DOCX successfully generated at {output_path}")
    except Exception as e:
        logger.error(f"Failed to generate DOCX: {e}", exc_info=True)
        raise e
